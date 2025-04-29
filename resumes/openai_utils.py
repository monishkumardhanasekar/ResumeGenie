import fitz
from docx import Document
import os
# from g4f.client import Client
from openai import OpenAI
from django.conf import settings
from pymongo import MongoClient
from dotenv import load_dotenv

load_dotenv()

openrouter_api_key = os.getenv("AI_API_KEY")


client = OpenAI(
  base_url="https://openrouter.ai/api/v1",
  api_key=openrouter_api_key
)

def extract_text_from_pdf(file):
    text = ""
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")  # Open PDF from file-like object
    for page in pdf_document:
        text += page.get_text()  # Extract text from each page
    return text

def extract_text_from_docx(file):
    text = ""
    try:
        doc = Document(file)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
    except Exception as e:
        print(f"An error occurred while extracting text from DOCX: {e}")
    return text

def extract_text_from_txt(file):
    text = file.read().decode('utf-8')  # Read the text file and decode it as UTF-8
    return text

def get_resume_details_from_ai(extracted_text):
    response = client.chat.completions.create(
        model="openai/gpt-3.5-turbo",
        messages=[
            {"role": "user", "content": f"Extract all details from the following resume text and provide a structured JSON format with fields like name, email, linkedin, phone, skills, work_experience, professional_certifications, education, project_experience and every fields with their details accordingly. \n\n{extracted_text}"}
        ],
        temperature=0.5,
        max_tokens=3000,
    )
    return response.choices[0].message.content
