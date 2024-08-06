import re
from datetime import datetime
import pdfplumber
from docx import Document
import zipfile
import xml.etree.ElementTree as ET
import fitz

def find_key_recursively(data, keys):
    if isinstance(data, dict):
        for k, v in data.items():
            if k in keys:
                return v
            result = find_key_recursively(v, keys)
            if result is not None:
                return result
    elif isinstance(data, list):
        for item in data:
            result = find_key_recursively(item, keys)
            if result is not None:
                return result
    return None


def check_contact_info(resume_json):
    score = 100
    weightage = {
        "name": 25,
        "email": 25,
        "phone": 25,
        "linkedin": 25
    }
    
    # Retrieve contact information
    name = find_key_recursively(resume_json, ['name'])
    email = find_key_recursively(resume_json, ['email'])
    phone = find_key_recursively(resume_json, ['phone'])
    linkedin = find_key_recursively(resume_json, ['linkedin'])
    
    # Check for each contact information field and deduct points if missing
    if not name:
        score -= weightage["name"]
    
    if not email:
        score -= weightage["email"]
    
    if not phone:
        score -= weightage["phone"]
    
    if not linkedin:
        score -= weightage["linkedin"]
    
    return score


def check_qualifications(resume_json):
    score = 100
    weightage = {
        "education": 30,
        "work_experience": 40,
        "professional_certifications": 30,
    }
    
    # Retrieve qualifications
    education = find_key_recursively(resume_json, ['education'])
    work_experience = find_key_recursively(resume_json, ['work_experience'])
    professional_certifications = find_key_recursively(resume_json, ['professional_certifications'])
    
    # Check for each qualification field and deduct points if missing
    if not education:
        score -= weightage["education"]
    
    if not work_experience:
        score -= weightage["work_experience"]
    
    if not professional_certifications:
        score -= weightage["professional_certifications"]
    
    return score


def check_education(resume_json):
    score = 100
    weightage = {
        "institutions_attended": 30,
        "degrees_earned": 40,
        "graduation_dates": 30,
    }
    
    # Define search keys
    education_keys = ['university', 'college', 'school', 'institute', 'institutions_attended', 'institution', 'academy', 'higher_education', 'academic_institution', 'learning_center', 'education', 'educational_institution', 'alumni', 'faculty', 'department', 'educational_center', 'technical_school', 'vocational_school', 'training_center', 'research_institute', 'graduate_school', 'postgraduate_institution', 'community_college', 'professional_school', 'educational_facility', 'study', 'educational_program', 'school_of_study', 'university_college', 'education_center']

    degree_keys = ['degree', 'qualification', 'credentials', 'degrees_earned', 'diploma', 'qualification', 'certification', 'degree_obtained', 'academic_degree', 'postgraduate_degree', 'bachelor', 'master', 'doctorate', 'PhD', 'undergraduate_degree', 'associate_degree', 'advanced_degree', 'professional_degree', 'diploma_certification', 'degree_certificate', 'educational_qualification', 'degree_title', 'degree_level', 'degree_awarded', 'honors_degree', 'specialization', 'major', 'minor','program']

    graduation_keys = ['dates', 'graduation_dates', 'expected_completion', 'start_date', 'end_date', 'graduation_year', 'completion_dates', 'culmination_dates', 'degree_conferred', 'duration', 'expected_graduation_date', 'expected_graduation', 'graduation_day', 'graduation_month', 'completion_date', 'date_of_graduation', 'completion_year', 'graduation_period', 'term_of_graduation', 'academic_completion', 'degree_completion', 'graduation_time', 'graduation_schedule', 'graduation_term', 'graduation_session', 'end_of_studies', 'completion_of_studies', 'degree_award_date', 'academic_award_date']

    
    # Retrieve education details using alternate keys under 'education'
    institutions_attended = find_key_recursively(resume_json.get('education', {}), education_keys)
    degrees_earned = find_key_recursively(resume_json.get('education', {}), degree_keys)
    graduation_dates = find_key_recursively(resume_json.get('education', {}), graduation_keys)
    
    # Check for each education field and deduct points if missing
    if not institutions_attended:
        score -= weightage["institutions_attended"]
    
    if not degrees_earned:
        score -= weightage["degrees_earned"]
    
    if not graduation_dates:
        score -= weightage["graduation_dates"]
    
    return score



def check_work_experience(resume_json):
    score = 100
    weightage = {
        "company_names": 20,
        "duration": 20,
        "location": 20,
        "title": 20,
        "responsibilities": 20,
        # "quantification_metrics": 20,
    }

    # Define search keys
    company_keys = ['company', 'employer', 'organization', 'company_names', 'workplace', 'business', 'firm', 'corporation', 'enterprise', 'institution', 'agency', 'office', 'entity', 'incorporated', 'co', 'limited', 'LLC', 'partnership', 'subsidiary', 'parent_company', 'association', 'group', 'outfit', 'venture', 'establishment']

    duration_keys = ['duration', 'time_period', 'dates', 'work_period', 'employment_dates', 'end_date', 'start_date', 'tenure', 'span', 'work_duration', 'years', 'months', 'period_of_employment', 'job_duration', 'employment_period', 'date_of_hiring', 'date_of_leaving', 'employment_span', 'contract_period', 'term', 'term_of_service', 'service_duration', 'service_period', 'employment_timeline', 'work_experience', 'job_timeline']

    location_keys = ['location', 'city', 'place', 'workplace_location', 'work_location', 'office_location', 'workplace_address', 'site', 'venue', 'region', 'area', 'district', 'locale', 'neighborhood', 'locality', 'address', 'zone', 'sector', 'province', 'state', 'territory', 'division', 'quarter', 'ward', 'department']

    title_keys = ['title', 'position', 'role', 'job_title', 'job_position', 'designation', 'occupation', 'function', 'rank', 'post', 'title_of_position', 'job_role', 'work_title', 'position_title', 'head', 'leader', 'manager', 'director', 'executive', 'chief', 'specialist', 'consultant', 'analyst', 'associate', 'officer', 'coordinator', 'supervisor', 'administrator', 'representative']

    responsibilities_keys = ['responsibilities', 'duties', 'tasks', 'job_responsibilities', 'work_responsibilities', 'obligations', 'functions', 'assignments', 'roles', 'activities', 'job_functions', 'work_duties', 'work_tasks', 'job_tasks', 'responsibilities_in_role', 'accountabilities', 'job_requirements', 'work_roles', 'work_expectations', 'job_experience','performance_criteria', 'work_requirements', 'areas_of_responsibility', 'job_description', 'work_description']



    # Retrieve work experience details using alternate keys under 'work_experience'
    company_names = find_key_recursively(resume_json.get('work_experience', {}), company_keys)
    duration = find_key_recursively(resume_json.get('work_experience', {}), duration_keys)
    location = find_key_recursively(resume_json.get('work_experience', {}), location_keys)
    title = find_key_recursively(resume_json.get('work_experience', {}), title_keys)
    responsibilities = find_key_recursively(resume_json.get('work_experience', {}), responsibilities_keys)

    # Check for each work experience field and deduct points if missing
    if not company_names:
        score -= weightage["company_names"]
    
    if not duration:
        score -= weightage["duration"]
    
    if not location:
        score -= weightage["location"]
    
    if not title:
        score -= weightage["title"]
    
    if not responsibilities:
        score -= weightage["responsibilities"]


    return score

def check_project_experience(resume_json):
    score = 80
    weightage = {
        "project_title": 40,
        "description": 40,
    }

    # Define search keys
    project_title_keys = ['name', 'names','project_title', 'title', 'project_name', 'project', 'projects', 'assignment', 'task_name', 'project_label', 'initiative', 'program_name', 'project_id', 'project_role', 'project_header', 'project_heading', 'project_topic', 'project_subject', 'project_identifier', 'project_name_label', 'project_title_description', 'project_code', 'project_tag', 'work_title', 'campaign_name', 'project_name_title', 'task_title', 'project_name_identifier', 'project_title_label', 'project_label_name', 'project_title_code', 'project_alias', 'project_reference']

    description_keys = ['description', 'project_description', 'details', 'project_details', 'summary', 'overview', 'narrative', 'project_summary', 'project_outline', 'explanation', 'report', 'analysis', 'project_report', 'description_of_project', 'project_explanation', 'project_specifications', 'project_summary_description', 'project_content', 'project_info', 'project_data', 'project_comments', 'project_notes', 'project_analysis', 'project_scope', 'project_statement', 'project_briefing', 'project_highlights', 'project_account', 'project_narrative', 'project_depiction','responsibilities']


    # Retrieve project experience details using alternate keys under 'project_experience'
    project_titles = find_key_recursively(resume_json.get('project_experience', {}), project_title_keys)
    descriptions = find_key_recursively(resume_json.get('project_experience', {}), description_keys)

    # Check for each project experience field and deduct points if missing
    if not project_titles:
        score -= weightage["project_title"]
    
    if not descriptions:
        score -= weightage["description"]

    return score



def check_quantification_metrics(resume_data):
    total_score = 20
    quantification_score = 20

    # Define regex patterns for quantification metrics
    metric_patterns = [
    r'\b\d+%?\b',  # Percentages and numbers
    r'\b\d+\s(?:people|employees|team|projects|tasks|clients|customers|users|sales|units)\b',  # Quantified units
    r'\b(?:sales|revenue|profit|budget|cost|efficiency|growth|output|performance|quality|capacity|retention|engagement)\b.*?\b\d+%?\b|\b\d+%?\b.*?\b(?:sales|revenue|profit|budget|cost|efficiency|growth|output|performance|quality|capacity|retention|engagement)\b',  # Performance metrics
    r'\b(?:\$?\d+\,?\d*\.?\d*\b|€\d+\,?\d*\.?\d*\b|£\d+\,?\d*\.?\d*\b)',  # Currency metrics
    r'\b(?:days|weeks|months|years|hours|minutes|seconds)\b.*?\b\d+%?\b|\b\d+%?\b.*?\b(?:days|weeks|months|years|hours|minutes|seconds)\b'  # Time metrics
    ]


    def has_quantification(text):
        for pattern in metric_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False

    all_jobs_have_quantification = True

    responsibilities_keys = ['responsibilities', 'duties', 'tasks', 'job_responsibilities', 'work_responsibilities', 'obligations', 'functions', 'assignments', 'roles', 'activities', 'job_functions', 'work_duties', 'work_tasks', 'job_tasks', 'responsibilities_in_role', 'accountabilities', 'job_requirements', 'work_roles', 'work_expectations', 'job_experience', 'performance_criteria', 'work_requirements', 'areas_of_responsibility', 'job_description', 'work_description']

    for job in resume_data.get('work_experience', []):
        job_has_quantification = False
        for key in responsibilities_keys:
            responsibilities = job.get(key, [])
            for responsibility in responsibilities:
                if has_quantification(responsibility):
                    job_has_quantification = True
                    break
            if job_has_quantification:
                break
        
        if not job_has_quantification:
            all_jobs_have_quantification = False
            break

    # Reduce score if any job title has no quantification metrics
    if not all_jobs_have_quantification:
        total_score -= quantification_score

    return total_score


import language_tool_python

def preprocess_text(text):
    # Remove email addresses
    text = re.sub(r'\S+@\S+', '', text)
    # Remove URLs
    text = re.sub(r'http\S+|www\S+', '', text)
    # Remove phone numbers
    text = re.sub(r'\+?\d[\d -]{8,12}\d', '', text)
    # Remove extra spaces
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def check_spelling_grammar(text):
    tool = language_tool_python.LanguageTool('en-US')  # Use 'en-US' for English (United States)

    # Preprocess the text
    cleaned_text = preprocess_text(text)

    matches = tool.check(cleaned_text)
    
    # Separate spelling and grammar errors
    spelling_errors = [match for match in matches if match.ruleId.startswith('MORFOLOGIK_RULE_EN_US')]
    grammar_errors = [match for match in matches if not match.ruleId.startswith('MORFOLOGIK_RULE_EN_US')]

    # Calculate scores
    total_score = 40
    spelling_score = 20
    grammar_score = 20

    if len(spelling_errors) > 10:
        spelling_score = 0  # Deduct full score if more than 10 spelling mistakes

    if len(grammar_errors) > 5:
        grammar_score = 0  # Deduct full score if more than 5 grammar mistakes

    total_score -= (20 - spelling_score)  # Deduct spelling score from total
    total_score -= (20 - grammar_score)   # Deduct grammar score from total

    # Prepare the results
    error_details = {
        'spelling': [],
        'grammar': []
    }

    for match in spelling_errors:
        error_details['spelling'].append({
            'message': match.message,
            'context': cleaned_text[match.offset:match.offset + match.errorLength],
            'replacements': ', '.join(r.value if hasattr(r, 'value') else r for r in match.replacements),
            'offset': match.offset,
            'error_length': match.errorLength
        })

    for match in grammar_errors:
        error_details['grammar'].append({
            'message': match.message,
            'context': cleaned_text[match.offset:match.offset + match.errorLength],
            'replacements': ', '.join(r.value if hasattr(r, 'value') else r for r in match.replacements),
            'offset': match.offset,
            'error_length': match.errorLength
        })

    return {
        'total_score': total_score,
        'spelling_score': spelling_score,
        'grammar_score': grammar_score,
        'error_details': error_details
    }


def check_pdf_for_images_and_tables(pdf_path):
    has_images = False
    has_tables = False

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            # Check for images
            if page.images:
                has_images = True
            
            # Check for tables
            tables = page.extract_tables()
            if tables:
                has_tables = True

            if has_images or has_tables:
                break  # Stop processing if either images or tables are found

    return has_images, has_tables

def calculate_pdf_score(pdf_path):
    has_images, has_tables = check_pdf_for_images_and_tables(pdf_path)
    score = 60
    if has_images:
        score -= 30
    if has_tables:
        score -= 30
    return score

def check_docx_for_elements(docx_path):
    doc = Document(docx_path)
    has_images = False
    has_tables = False
    has_graphs = False

    # Check for images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            has_images = True
            break

    # Check for tables
    if doc.tables:
        has_tables = True

    # Check for charts using zipfile and xml.etree.ElementTree
    with zipfile.ZipFile(docx_path, 'r') as docx:
        for file in docx.namelist():
            if file.startswith('word/charts/'):
                with docx.open(file) as chart_file:
                    tree = ET.parse(chart_file)
                    root = tree.getroot()
                    if root.tag.endswith('chart'):
                        has_graphs = True
                        break

    return has_images, has_tables, has_graphs

def calculate_docx_score(docx_path):
    has_images, has_tables, has_graphs = check_docx_for_elements(docx_path)
    score = 60
    if has_images:
        score -= 20
    if has_tables:
        score -= 20
    if has_graphs:
        score -= 20
    return score, has_images, has_tables, has_graphs


# # to get the list of the fonts from docx
# def get_docx_fonts(file_path):
#     doc = Document(file_path)
#     fonts = set()

#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             if run.font.name:
#                 fonts.add(run.font.name)

#     return fonts


# # to get the list of the fonts from pdf
# def get_pdf_fonts(file):
#     # Move the stream position to the beginning
#     file.seek(0)
#     pdf_document = fitz.open(stream=file.read(), filetype="pdf")
#     fonts = set()
    
#     # Iterate over each page in the PDF
#     for page_num in range(pdf_document.page_count):
#         page = pdf_document[page_num]
#         # Get text information for the page
#         text_page = page.get_textpage()
#         # Extract font information
#         for block in text_page.extractDICT()["blocks"]:
#             for line in block["lines"]:
#                 for span in line["spans"]:
#                     fonts.add(span["font"])
    
#     pdf_document.close()
#     return list(fonts)



ALLOWED_FONTS = {
    'Times New Roman', 'TimesNewRomanPS-ItalicMT', 'TimesNewRomanPS-BoldMT', 'TimesNewRomanPS-BoldItal', 'TimesNewRomanPSMT',
    'Tahoma', 'Tahoma-ItalicMT', 'Tahoma-BoldMT', 'Tahoma-BoldItal',
    'Verdana', 'Verdana-ItalicMT', 'Verdana-BoldMT', 'Verdana-BoldItal',
    'Arial', 'Arial-BoldMT', 'Arial-ItalicMT', 'Arial-BoldItalicMT', 'ArialMT',
    'Helvetica', 'Helvetica-ItalicMT', 'Helvetica-BoldMT', 'Helvetica-BoldItalicMT',
    'Calibri', 'Calibri-ItalicMT', 'Calibri-BoldMT', 'Calibri-BoldItalicMT',
    'Georgia', 'Georgia-ItalicMT', 'Georgia-BoldMT', 'Georgia-BoldItalicMT',
    'Cambria', 'Cambria-ItalicMT', 'Cambria-BoldMT', 'Cambria-BoldItalicMT',
    'Gill Sans', 'GillSans-ItalicMT', 'GillSans-BoldMT', 'GillSans-BoldItalicMT',
    'Garamond', 'Garamond-ItalicMT', 'Garamond-BoldMT', 'Garamond-BoldItalicMT'
}

MIN_FONT_SIZE = 9
MAX_FONT_SIZE = 16
FONT_CHECK_SCORE = 20
SIZE_CHECK_SCORE = 20
EMPTY_FONT_PENALTY = 20  # Penalty for not finding any fonts
EMPTY_SIZE_PENALTY = 20  # Penalty for not finding any sizes

def get_pdf_fonts(file):
    file.seek(0)
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    fonts = set()
    sizes = set()
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text_page = page.get_textpage()
        for block in text_page.extractDICT()["blocks"]:
            for line in block["lines"]:
                for span in line["spans"]:
                    fonts.add(span["font"])
                    sizes.add(span["size"])

    return fonts, sizes

def get_docx_fonts(file):
    document = Document(file)
    fonts = set()
    sizes = set()

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.font.name:
                fonts.add(run.font.name)
            if run.font.size:
                sizes.add(run.font.size.pt)

    return fonts, sizes

def calculate_pdf_fonts_score(file):
    fonts, sizes = get_pdf_fonts(file)
    
    # Calculate font check score
    font_check_score = FONT_CHECK_SCORE
    if not fonts:
        font_check_score -= EMPTY_FONT_PENALTY
    else:
        for font in fonts:
            if font not in ALLOWED_FONTS:
                font_check_score -= FONT_CHECK_SCORE
                break

    # Calculate size check score
    size_check_score = SIZE_CHECK_SCORE
    if not sizes:
        size_check_score -= EMPTY_SIZE_PENALTY
    else:
        for size in sizes:
            if size < MIN_FONT_SIZE or size > MAX_FONT_SIZE:
                size_check_score -= SIZE_CHECK_SCORE
                break

    return font_check_score + size_check_score

def calculate_docx_fonts_score(file):
    fonts, sizes = get_docx_fonts(file)
    
    # Calculate font check score
    font_check_score = FONT_CHECK_SCORE
    if not fonts:
        font_check_score -= EMPTY_FONT_PENALTY
    else:
        for font in fonts:
            if font not in ALLOWED_FONTS:
                font_check_score -= FONT_CHECK_SCORE
                break

    # Calculate size check score
    size_check_score = SIZE_CHECK_SCORE
    if not sizes:
        size_check_score -= EMPTY_SIZE_PENALTY
    else:
        for size in sizes:
            if size < MIN_FONT_SIZE or size > MAX_FONT_SIZE:
                size_check_score -= SIZE_CHECK_SCORE
                break

    return font_check_score + size_check_score
